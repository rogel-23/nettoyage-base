import pandas as pd
import numpy as np
from docx import Document
import unidecode
import re
import streamlit as st
from sklearn.neighbors import NearestNeighbors
# from sentence_transformers import SentenceTransformer
import unidecode
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import openpyxl
from openpyxl.styles import PatternFill
from datetime import datetime
from constants import (
    abbreviation_mapping, id_columns, date_columns, entry_date_columns, date_exceptions,
    entry_year_exceptions, seniority_date_columns, seniority_exceptions, position_entry_date_columns,
    exit_date_columns, gender_columns, hours_columns, workload_columns, workingtime_columns,
    contract_columns, status_columns, region_columns, postal_code_columns, city_columns, contrat_mapping,
    segmentation_columns, job_title_columns, job_columns, nplus1_columns, convention_columns, base_salary_columns,
    coefficient_columns, idcc_columns
)
# from utils_erreurs import log_error, generate_word_report


# Rapport Word
def generate_word_report(report_type, content):
    doc = Document()
    doc.add_heading(f"Rapport {report_type}", level=1)
    for item in content:
        doc.add_paragraph(item)
    filename = f"rapport_{report_type.lower().replace(' ', '_')}.docx"
    doc.save(filename)
    return filename

def is_valid_postal(val, empty_check=True):
    if pd.isna(val) or str(val).strip() == "":
        return not empty_check  # False = erreur si vide et case cochée
    return bool(re.fullmatch(r"\d{5}", str(val).strip()))


def harmonize_text(value, abbreviation_dict=None):
    if pd.isna(value):
        return value

    text = str(value)
    text = unidecode.unidecode(text)  # Supprime les accents
    text = text.upper().strip()       # Majuscules et retrait espaces
    text = re.sub(r"\s+", " ", text)  # Supprime les doubles espaces 

    # Remplacement des abréviations connues
    if abbreviation_dict:
        for abbr, full in abbreviation_dict.items():
            # Remplace uniquement les mots entiers (pas au milieu d'autres mots)
            pattern = r"\b" + re.escape(abbr) + r"\b"
            text = re.sub(pattern, full, text)

        return text

def harmonize_if_known(val, contrat_mapping):
    if pd.isna(val) or str(val).strip() == "":
        return val  # on garde la valeur telle quelle
    key = str(val).strip().lower()
    return contrat_mapping.get(key, val)  # si trouvé : valeur harmonisée, sinon : original

# Identifier les erreurs pour le rapport
def is_invalid_contract(val, contrat_mapping=None, empty_check=True):
    if pd.isna(val) or str(val).strip() == "":
        return empty_check
    key = str(val).strip().lower()
    return key not in contrat_mapping

def clean_workload(val):
    if pd.isna(val) or str(val).strip() == "":
        return np.nan

    try:
        val_str = str(val).replace(",", ".").replace("%", "").strip()
        val_float = float(val_str)

        # Si c'est une proportion (inférieure à 1), on convertit
        if 0 < val_float <= 1:
            val_float *= 100

        val_float = round(val_float, 1)

        if val_float.is_integer():
            return f"{int(val_float)}%"
        else:
            return f"{val_float:.1f}%"

    except:
        return "Non reconnu"
    
# Erreurs : valeurs < 0 ou > 100 ou "Non reconnu"
def is_invalid_workload(val, empty_check=True):
    if pd.isna(val) or str(val).strip() == "":
        return empty_check  # True = erreur si vide et case cochée

    if val == "Non reconnu":
        return True

    try:
        val_float = float(str(val).replace(",", ".").replace("%", "").strip())
        
        # Même logique que clean_workload
        if 0 < val_float <= 1:
            val_float *= 100
        
        return val_float < 0 or val_float > 100

    except:
        return True

    
def extract_hour(val):
    if pd.isna(val) or str(val).strip() == "":
        return np.nan  # Case vide → pas d'erreur
    try:
        val = str(val).lower().replace(",", ".")
        match = re.search(r"(\d+(?:\.\d+)?)", val)
        return round(float(match.group(1)), 1) if match else "Non reconnu"
    except:
        return "Non reconnu"
    
# Initialisation du modèle
# model = SentenceTransformer("sentence-transformers/paraphrase-xlm-r-multilingual-v1")

def normalize_job_title(title):
    if pd.isna(title):
        return ""
    title = str(title)
    title = unidecode.unidecode(title).upper()
    title = re.sub(r"[^\w\s]", "", title)
    title = re.sub(r"\s+", " ", title.strip())
    return title

'''def suggest_job_title_mapping(series, threshold=0.85):
    cleaned_series = series.fillna("").apply(normalize_job_title)
    unique_titles = cleaned_series.unique().tolist()

    embeddings = model.encode(unique_titles, convert_to_tensor=True)
    similarity_matrix = cosine_similarity(embeddings)

    groups = {}
    used = set()

    for i, title in enumerate(unique_titles):
        if title in used:
            continue
        group = [title]
        for j in range(i + 1, len(unique_titles)):
            if similarity_matrix[i][j] >= threshold:
                group.append(unique_titles[j])
                used.add(unique_titles[j])
        used.add(title)
        group_leader = sorted(group, key=len)[0]
        for t in group:
            groups[t] = group_leader

    mapping_df = pd.DataFrame(list(groups.items()), columns=["Original", "Suggestion"])
    mapping_df["Harmonisation finale"] = mapping_df["Suggestion"]
    initial_mapping = dict(zip(mapping_df["Original"], mapping_df["Suggestion"]))

    return mapping_df, initial_mapping'''

def harmonize_financial_values(df, col, empty_check=True):
    original = df[col].copy()
    error_rows = []

    for i, val in enumerate(original):
        str_val = str(val).replace(",", ".").strip()

        if str_val == "" or pd.isna(val):
            if empty_check:
                error_rows.append((i, val))
            continue

        try:
            num = float(str_val)
            if num < 0:
                error_rows.append((i, val))
            else:
                df.at[i, col] = num  # ✅ Conversion uniquement si valide
        except:
            error_rows.append((i, val))  # Laisser la valeur telle quelle

    modified = [(idx, col) for idx, _ in error_rows]
    errors = [f"{col} : {len(error_rows)} valeur(s) négative(s), vide(s) ou invalide(s)."] if error_rows else []

    return df, errors, modified


'''def get_similar_job_title_groups(series, threshold=0.85):
    series = series.fillna("").apply(normalize_job_title)
    unique_titles = series.unique().tolist()
    embeddings = model.encode(unique_titles, convert_to_tensor=True)
    similarity_matrix = cosine_similarity(embeddings)

    groups = []
    used = set()

    for i, title in enumerate(unique_titles):
        if title in used:
            continue
        group = [title]
        for j in range(i + 1, len(unique_titles)):
            if similarity_matrix[i][j] >= threshold and unique_titles[j] not in used:
                group.append(unique_titles[j])
                used.add(unique_titles[j])
        if len(group) > 1:
            groups.append(group)
            used.update(group)

    return series, groups'''


def save_cleaned_excel(df, modified_cells, columns_to_check_dupes, incoherent_entry_dates, column_menus, smic_threshold):

    red_fill = PatternFill(start_color="FF9999", end_color="FF9999", fill_type="solid")
    date_pattern = re.compile(r"^\d{2}/\d{2}/\d{4}$")
    clean_filename = "données_nettoyées.xlsx"

    df.replace({pd.NaT: None, np.nan: None}, inplace=True)
    df.to_excel(clean_filename, index=False, engine="openpyxl")

    wb = openpyxl.load_workbook(clean_filename)
    ws = wb.active

    idcc_errors = st.session_state.get("invalid_idcc", {})
    incoherent_dates_map = st.session_state.get("incoherent_entry_dates", {})

    for key, val in incoherent_dates_map.items():
        if isinstance(val, set):
            incoherent_dates_map[key] = list(val)

    dupes_map = {
        col: set(df[df.duplicated(subset=[col], keep=False)].index.tolist())
        for col in columns_to_check_dupes
    }

    def check_empty_cell_error(col_name, val):
        default = False if col_name in idcc_columns else True
        return st.session_state.get(f"check_empty_{col_name}", default) and (val is None or str(val).strip() == "")

    for row_idx, row in enumerate(ws.iter_rows(min_row=2, max_row=ws.max_row), start=2):
        for col_idx, col_name in enumerate(df.columns, start=1):
            cell = ws.cell(row=row_idx, column=col_idx)
            val = cell.value
            is_error = False

            if check_empty_cell_error(col_name, val):
                is_error = True

            if col_name in gender_columns and val == "Non reconnu":
                is_error = True

            if col_name in idcc_errors and (row_idx - 2) in idcc_errors[col_name]:
                val_is_empty = val is None or str(val).strip() == ""
                empty_check = st.session_state.get(f"check_empty_{col_name}", False)
                if not val_is_empty or (val_is_empty and empty_check):
                    is_error = True


            if col_name in nplus1_columns:
                conflit_col = st.session_state.get("_nplus1_conflit_")
                if conflit_col is not None:
                    conflit_val = conflit_col.iloc[row_idx - 2]
                    if isinstance(conflit_val, (bool, np.bool_)) and conflit_val:
                        is_error = True

            if col_name in postal_code_columns and column_menus.get(col_name) == "📮 Vérification code postal":
                if not is_error:  # éviter de surcharger l'erreur vide
                    is_error = not re.fullmatch(r"\d{5}", str(val).strip()) if val else False
                cell.number_format = "@"

            if col_name in st.session_state.get("invalid_gender", {}) and (row_idx - 2) in st.session_state["invalid_gender"][col_name]:
                is_error = True

            if col_name in st.session_state.get("invalid_workload", {}) and (row_idx - 2) in st.session_state["invalid_workload"][col_name]:
                is_error = True

            if col_name in hours_columns and val == "Non reconnu":
                is_error = True

            if column_menus.get(col_name) == "💶 Vérification salaire de base":
                if val is None or str(val).strip() == "":
                    if st.session_state.get(f"check_empty_{col_name}", True):
                        is_error = True
                else:
                    try:
                        val_float = float(str(val).replace(",", ".").strip())
                        if val_float < 0 or val_float < smic_threshold:
                            is_error = True 
                    except:
                        is_error = True

            if column_menus.get(col_name) == "💰 Harmonisation des valeurs financières":
                try:
                    if val is not None and str(val).strip() != "":
                        val_float = float(str(val).replace(",", ".").strip())
                        if val_float < 0:
                            is_error = True
                except:
                    is_error = True

            if col_name in contract_columns:
                if val and str(val).strip().lower() not in contrat_mapping:
                    is_error = True

            if col_name in incoherent_dates_map and (row_idx - 2) in incoherent_dates_map[col_name]:
                is_error = True
                cell.number_format = "DD/MM/YYYY"

            if col_name in st.session_state.get("invalid_dates", {}) and (row_idx - 2) in st.session_state["invalid_dates"][col_name]:
                is_error = True

            if col_name in st.session_state.get("invalid_hours", {}) and (row_idx - 2) in st.session_state["invalid_hours"][col_name]:
                is_error = True

            if col_name in dupes_map and (row_idx - 2) in dupes_map[col_name]:
                is_error = True

            if col_name in coefficient_columns:
                coeff_errors = st.session_state.get("coeff_format_errors", {}).get(col_name, [])
                if (row_idx - 2) in coeff_errors:
                    is_error = True

            if col_name in date_columns and column_menus.get(col_name) == "📆 Harmoniser les dates":
                if isinstance(val, datetime):
                    if val.date() == datetime(1900, 1, 1).date():
                        is_error = True
                    cell.number_format = "DD/MM/YYYY"
                elif isinstance(val, str):
                    try:
                        parsed = pd.to_datetime(val, errors="raise", dayfirst=True)
                        cell.value = parsed.strftime("%d/%m/%Y")
                        cell.number_format = "DD/MM/YYYY"
                    except:
                        is_error = True

            if is_error:
                cell.fill = red_fill

    wb.save(clean_filename)
    # 🔧 Ajustement automatique de la largeur des colonnes
    for column_cells in ws.columns:
        max_length = 0
        column = column_cells[0].column_letter  # ex: "A", "B"...
        for cell in column_cells:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = max_length + 2  # petit padding
        ws.column_dimensions[column].width = adjusted_width

    wb.close()

    with open(clean_filename, "rb") as f:
        st.session_state.excel_bytes = f.read()
    st.session_state.ready_to_download = True

    if "_nplus1_conflit_" in df.columns:
        df.drop(columns=["_nplus1_conflit_"], inplace=True)


def apply_fusions(df, column_menus, fusion_mode=False):
    modified_cells = set()
    modifications = []

    for col, action in column_menus.items():
        if action in ["👔 Harmoniser les intitulés d'emploi", "🤝 Harmoniser la convention collective"]:
            # 1. Harmonisation textuelle systématique
            original = df[col].copy()
            df[col] = df[col].fillna("").apply(normalize_job_title)

            changed_text = df[col] != original
            if changed_text.any():
                modifications.append(f"Harmonisation textuelle appliquée sur {col}")
                modified_cells.update([(idx, col) for idx in df[changed_text].index])

            # 2. PAS DE FUSION — suppression du bloc fusion_mapping

    return df, modifications, modified_cells


def verify_base_salary(df, col, smic_threshold, empty_check=True):
    original = df[col].copy()
    error_rows = []

    for i, val in enumerate(original):
        str_val = str(val).replace(",", ".").strip()

        if str_val == "" or pd.isna(val):
            if empty_check:
                error_rows.append((i, val))
            continue

        try:
            num = float(str_val)
            if num < 0 or num < smic_threshold:
                error_rows.append((i, val))
            else:
                df.at[i, col] = num
        except:
            error_rows.append((i, val))

    # Ne pas forcer le type ici — on laisse Pandas gérer la mixité
    errors = [f"{col} : {len(error_rows)} valeur(s) < SMIC, invalide(s) ou vide(s)."] if error_rows else []
    modified = [(idx, col) for idx, _ in error_rows]

    return df, errors, modified
